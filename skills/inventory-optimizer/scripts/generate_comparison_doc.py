#!/usr/bin/env python3
"""
Generate branded Word document with vendor price comparison tables.
Uses python-docx to create professional BCEHS-branded comparison documents.
"""

import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# BCEHS Brand Colors
BCEHS_BLUE = RGBColor(0, 51, 102)
BCEHS_LIGHT_BLUE = RGBColor(91, 155, 213)
GRAY = RGBColor(127, 127, 127)
LIGHT_GRAY = RGBColor(217, 217, 217)
WHITE = RGBColor(255, 255, 255)


def set_cell_background(cell, color):
    """Set table cell background color."""
    shading_elm = OxmlElement('w:shd')
    # Convert RGBColor to hex string
    hex_color = f'{color[0]:02x}{color[1]:02x}{color[2]:02x}'
    shading_elm.set(qn('w:fill'), hex_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_comparison_document(comparison_data, output_path):
    """
    Create a Word document with vendor price comparison.
    
    Args:
        comparison_data: Dict with structure:
            {
                "product_name": "Branded Pens",
                "date": "2026-01-30",
                "vendors": [
                    {
                        "name": "Creative Boulevard",
                        "pricing": [
                            {"quantity": 100, "total": 250.00, "unit_cost": 2.50, "shipping": 15.00},
                            {"quantity": 500, "total": 1000.00, "unit_cost": 2.00, "shipping": 25.00}
                        ],
                        "notes": "Setup fee: $50"
                    }
                ]
            }
        output_path: Path to save the .docx file
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Vendor Price Comparison', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = BCEHS_BLUE
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    
    # Subtitle with product and date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_text = subtitle.add_run(
        f"{comparison_data.get('product_name', 'Product')}\n"
        f"Prepared for: BCEHS Talent Acquisition\n"
        f"Date: {comparison_data.get('date', datetime.now().strftime('%Y-%m-%d'))}"
    )
    subtitle_text.font.size = Pt(11)
    subtitle_text.font.color.rgb = GRAY
    
    doc.add_paragraph()  # Spacing
    
    # For each vendor, create a section with pricing table
    vendors = comparison_data.get('vendors', [])
    
    for i, vendor in enumerate(vendors):
        if i > 0:
            doc.add_page_break()
        
        # Vendor name header
        vendor_header = doc.add_heading(vendor['name'], level=1)
        vendor_header_run = vendor_header.runs[0]
        vendor_header_run.font.color.rgb = BCEHS_BLUE
        vendor_header_run.font.size = Pt(14)
        
        # Pricing table
        pricing = vendor.get('pricing', [])
        if pricing:
            # Create table with header row
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Quantity', 'Unit Cost', 'Subtotal', 'Shipping', 'Total Cost']
            
            for idx, header_text in enumerate(headers):
                cell = header_cells[idx]
                cell.text = header_text
                set_cell_background(cell, BCEHS_BLUE)
                
                # Format header text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = WHITE
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for idx, price_point in enumerate(pricing):
                row_cells = table.add_row().cells
                
                qty = price_point.get('quantity', 0)
                unit = price_point.get('unit_cost', 0)
                shipping = price_point.get('shipping', 0)
                subtotal = qty * unit
                total = subtotal + shipping
                
                row_cells[0].text = f"{qty:,}"
                row_cells[1].text = f"${unit:.2f}"
                row_cells[2].text = f"${subtotal:,.2f}"
                row_cells[3].text = f"${shipping:.2f}"
                row_cells[4].text = f"${total:,.2f}"
                
                # Alternate row colors
                bg_color = LIGHT_GRAY if idx % 2 == 1 else WHITE
                for cell in row_cells:
                    set_cell_background(cell, bg_color)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Vendor notes
        notes = vendor.get('notes')
        if notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph()
            notes_run = notes_para.add_run(f"Notes: {notes}")
            notes_run.font.size = Pt(10)
            notes_run.font.italic = True
            notes_run.font.color.rgb = GRAY
    
    # Summary section
    doc.add_page_break()
    summary_header = doc.add_heading('Recommendation Summary', level=1)
    summary_header_run = summary_header.runs[0]
    summary_header_run.font.color.rgb = BCEHS_BLUE
    summary_header_run.font.size = Pt(14)
    
    # Find best vendor at each common quantity
    quantities = set()
    for vendor in vendors:
        for price_point in vendor.get('pricing', []):
            quantities.add(price_point.get('quantity'))
    
    if quantities:
        summary_para = doc.add_paragraph()
        summary_para.add_run("Best pricing by quantity:\n\n").bold = True
        
        for qty in sorted(quantities):
            best_vendor = None
            best_price = float('inf')
            
            for vendor in vendors:
                for price_point in vendor.get('pricing', []):
                    if price_point.get('quantity') == qty:
                        total = (qty * price_point.get('unit_cost', 0)) + price_point.get('shipping', 0)
                        if total < best_price:
                            best_price = total
                            best_vendor = vendor['name']
            
            if best_vendor:
                summary_para.add_run(f"• {qty:,} units: {best_vendor} (${best_price:,.2f} total)\n")
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "This comparison is for internal review purposes. "
        "Final vendor selection should consider quality, lead time, and other factors beyond price."
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = GRAY
    
    # Save document
    doc.save(output_path)
    return output_path


def main():
    """CLI entry point."""
    if len(sys.argv) < 3:
        print("Usage: python generate_comparison_doc.py <json_file> <output_path>")
        print("\nJSON structure:")
        print(json.dumps({
            "product_name": "Branded Pens",
            "date": "2026-01-30",
            "vendors": [
                {
                    "name": "Vendor A",
                    "pricing": [
                        {"quantity": 100, "unit_cost": 2.50, "shipping": 15.00}
                    ],
                    "notes": "Optional notes"
                }
            ]
        }, indent=2))
        sys.exit(1)
    
    json_file = sys.argv[1]
    output_path = sys.argv[2]
    
    with open(json_file, 'r') as f:
        comparison_data = json.load(f)
    
    result = create_comparison_document(comparison_data, output_path)
    print(f"✓ Comparison document created: {result}")


if __name__ == '__main__':
    main()
